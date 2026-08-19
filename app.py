import streamlit as st
from pathlib import Path
import pandas as pd, re, csv, hashlib, tempfile, shutil, subprocess, zipfile, datetime, os, traceback
from io import BytesIO

st.set_page_config(page_title="GEU AQAR Consolidator",page_icon="📊",layout="wide")
SUPPORTED={".xlsx",".xls",".csv",".docx",".pdf",".txt"}
EXCLUDE={"output","__pycache__",".git",".venv","venv"}

@st.cache_data
def rules():
    with open(Path(__file__).parent/"metric_rules.csv",encoding="utf-8-sig") as f:return list(csv.DictReader(f))
@st.cache_data
def nt_rules():
    with open(Path(__file__).parent/"non_template_requirements.csv",encoding="utf-8-sig") as f:return {x["metric"]:x["required_information"] for x in csv.DictReader(f)}
RULES=rules(); NT=nt_rules()

def excluded(p):
    return any(x in EXCLUDE for x in p.parts) or p.name.startswith(("~$",".~lock.","AQAR_Master_Consolidation"))

def metric(text):
    for r in RULES:
        if re.search(r"(?<!\d)"+re.escape(r["metric"])+r"(?!\d)",str(text),re.I):return r["metric"]
    return ""

def dept(root,p):
    q=p.relative_to(root); return q.parts[0] if len(q.parts)>1 else root.name

def criterion(p):
    for x in p.parts:
        m=re.match(r"(?i)criteria[-_ ]?(\d+)",x)
        if m:return "Criterion "+m.group(1)
    return ""

def hashfile(p):
    h=hashlib.sha256()
    try:
        with open(p,"rb") as f:
            for b in iter(lambda:f.read(1048576),b):h.update(b)
    except:return ""
    return h.hexdigest()

def lo_convert(p):
    exe=shutil.which("libreoffice") or shutil.which("soffice")
    if not exe:return None
    td=Path(tempfile.mkdtemp(prefix="aqarlo_"))
    try:
        subprocess.run([exe,"--headless","--convert-to","xlsx","--outdir",str(td),str(p)],stdout=subprocess.PIPE,stderr=subprocess.PIPE,timeout=120)
        q=td/(p.stem+".xlsx"); return q if q.exists() else None
    except:return None

def excel(p):
    try:return pd.ExcelFile(p),None
    except Exception as e:
        q=lo_convert(p)
        if q:
            try:return pd.ExcelFile(q),f"LibreOffice fallback used: {e}"
            except Exception as e2:return None,f"Excel and fallback failed: {e}; {e2}"
        return None,str(e)

def scan(root,cb):
    files=[p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED and not excluded(p)]
    inv=[]; rec=[]; errors=[]
    for i,p in enumerate(files,1):
        if cb:cb(i,len(files),p)
        d=dept(root,p); c=criterion(p); h=hashfile(p)
        inv.append({"Department":d,"Criterion":c,"Relative Path":str(p.relative_to(root)),"File":p.name,"Extension":p.suffix.lower(),"Size":p.stat().st_size,"SHA256":h})
        try:
            if p.suffix.lower() in {".xlsx",".xls"}:
                xl,note=excel(p)
                if not xl:errors.append({"File":str(p),"Error":note});continue
                if note:errors.append({"File":str(p),"Error":note})
                for sh in xl.sheet_names:
                    try:df=pd.read_excel(xl,sheet_name=sh,header=None)
                    except Exception as e:errors.append({"File":str(p),"Error":f"{sh}: {e}"});continue
                    for i,row in df.iterrows():
                        vals=[str(x) for x in row.tolist() if pd.notna(x)]
                        m=metric(sh+" "+" | ".join(vals))
                        if m:
                            rr=next(x for x in RULES if x["metric"]==m)
                            rec.append({"AQAR Metric":m,"Department":d,"Criterion":c,"Information Stream":"Data Template" if rr["non_template"]=="NO" else "Both / Review","File":p.name,"Relative Path":str(p.relative_to(root)),"Sheet":sh,"Row":i+1,"Content":" | ".join(vals[:30]),"File Hash":h})
            else:
                text=""
                if p.suffix.lower()==".txt":text=p.read_text(errors="ignore")
                elif p.suffix.lower()==".pdf":
                    import pypdf;text="\n".join((x.extract_text() or "") for x in pypdf.PdfReader(str(p)).pages)
                elif p.suffix.lower()==".docx":
                    from docx import Document
                    x=Document(str(p));text="\n".join(z.text for z in x.paragraphs)
                    text+="\n"+"\n".join(" | ".join(cell.text for cell in row.cells) for t in x.tables for row in t.rows)
                m=metric(p.name+" "+text)
                if m:rec.append({"AQAR Metric":m,"Department":d,"Criterion":c,"Information Stream":"Information Not in Data Template" if m in NT else "Supporting File","File":p.name,"Relative Path":str(p.relative_to(root)),"Sheet":"","Row":"","Content":text[:12000],"File Hash":h})
        except Exception as e:errors.append({"File":str(p),"Error":traceback.format_exc()})
    return files,pd.DataFrame(inv),pd.DataFrame(rec),pd.DataFrame(errors)

def workbook(root,inv,raw,err):
    if raw.empty:raw=pd.DataFrame(columns=["AQAR Metric","Department","Criterion","Information Stream","File","Relative Path","Sheet","Row","Content","File Hash"])
    deps=sorted(inv["Department"].dropna().unique()) if not inv.empty else []
    props=[]
    for m,g in raw.groupby("AQAR Metric"):
        r=next(x for x in RULES if x["metric"]==m)
        props.append({"AQAR Metric":m,"Title":r["title"],"Aggregation Rule":r["aggregation"],"Departments":g.Department.nunique(),"Records":len(g),"Status":"PENDING IQAC APPROVAL"})
    evidence=[]
    for m,req in NT.items():
        for d in deps:
            g=raw[(raw["AQAR Metric"]==m)&(raw.Department==d)]
            evidence.append({"Department":d,"AQAR Metric":m,"Required Information":req,"Submitted":"YES" if len(g) else "NO","Files":"; ".join(g["Relative Path"].unique()) if len(g) else "","IQAC Status":"VERIFY" if len(g) else "MISSING - ACTION REQUIRED"})
    missing=[{"AQAR Metric":r["metric"],"Title":r["title"],"Non-template":r["non_template"],"Found":r["metric"] in set(raw["AQAR Metric"])} for r in RULES]
    dup=raw[raw.duplicated(["AQAR Metric","Content"],keep=False)] if len(raw) else raw
    conflicts=[{"AQAR Metric":m,"Departments":g.Department.nunique(),"Records":len(g),"Status":"REVIEW REQUIRED"} for m,g in raw.groupby("AQAR Metric") if len(g)>1]
    b=BytesIO()
    with pd.ExcelWriter(b,engine="openpyxl") as w:
        pd.DataFrame([{"AQAR Root":str(root),"Files":len(inv),"Records":len(raw),"Metrics":raw["AQAR Metric"].nunique(),"Departments":len(deps),"Generated":str(datetime.datetime.now())}]).to_excel(w,index=False,sheet_name="ReadMe")
        inv.to_excel(w,index=False,sheet_name="File Inventory");pd.DataFrame(RULES).to_excel(w,index=False,sheet_name="Metric Rules");raw.to_excel(w,index=False,sheet_name="Raw Extract")
        pd.DataFrame(props).to_excel(w,index=False,sheet_name="Proposed Consolidation");dup.to_excel(w,index=False,sheet_name="Duplicate Check");pd.DataFrame(conflicts).to_excel(w,index=False,sheet_name="Conflicts")
        pd.DataFrame(missing).to_excel(w,index=False,sheet_name="Missing Metrics");pd.DataFrame(evidence).to_excel(w,index=False,sheet_name="Information Not in Data Template")
        raw.drop_duplicates().to_excel(w,index=False,sheet_name="Evidence Mapping");err.to_excel(w,index=False,sheet_name="Errors")
    b.seek(0);return b

st.title("📊 GEU AQAR Consolidator")
st.caption("Local OneDrive folder → AQAR metric extraction → non-template evidence → IQAC review")
mode=st.sidebar.radio("Source",["Local OneDrive / folder","Upload ZIP"])
if mode=="Local OneDrive / folder":
    path=st.text_input("Local AQAR root folder",placeholder="/Users/yourname/OneDrive/AQAR")
    if path:
        root=Path(path).expanduser()
        if root.is_dir():
            fs=[p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED and not excluded(p)]
            ds=sorted(set(dept(root,p) for p in fs))
            a,b,c=st.columns(3);a.metric("Files",len(fs));b.metric("Departments",len(ds));c.metric("Non-template metrics",len(NT))
            st.write("Departments:",", ".join(ds) if ds else "None")
            if st.button("🚀 Scan and Consolidate",type="primary"):
                bar=st.progress(0);msg=st.empty()
                def cb(i,n,p):bar.progress(i/n if n else 1);msg.write(f"Scanning {i:,}/{n:,}: `{p.name}`")
                with st.spinner("Scanning..."):files,inv,raw,err=scan(root,cb)
                st.session_state.result=(root,inv,raw,err);st.success(f"Completed: {len(files):,} files, {len(raw):,} records, {len(err):,} errors.")
        else:st.error("Folder does not exist.")
else:
    up=st.file_uploader("Upload ZIP (use only for small tests; do not upload your 7 GB repository)",type=["zip"])
    if up and st.button("🚀 Scan ZIP",type="primary"):
        td=Path(tempfile.mkdtemp(prefix="aqarzip_"));zp=td/up.name;zp.write_bytes(up.getbuffer())
        with zipfile.ZipFile(zp) as z:z.extractall(td/"x")
        root=td/"x"; dirs=[p for p in root.rglob("*") if p.is_dir() and any(x.is_file() and x.suffix.lower() in SUPPORTED for x in p.rglob("*"))]
        if dirs:root=max(dirs,key=lambda p:sum(1 for x in p.rglob("*") if x.is_file() and x.suffix.lower() in SUPPORTED))
        files,inv,raw,err=scan(root,None);st.session_state.result=(root,inv,raw,err);st.success(f"Completed: {len(files):,} files, {len(raw):,} records, {len(err):,} errors.")
if "result" in st.session_state:
    root,inv,raw,err=st.session_state.result
    a,b,c,d=st.columns(4);a.metric("Files",len(inv));b.metric("Records",len(raw));c.metric("Metrics",raw["AQAR Metric"].nunique() if len(raw) else 0);d.metric("Errors",len(err))
    t1,t2,t3=st.tabs(["Proposed Consolidation","Information Not in Data Template","Evidence / Errors"])
    props=[]
    for m,g in raw.groupby("AQAR Metric"):
        r=next(x for x in RULES if x["metric"]==m);props.append({"AQAR Metric":m,"Title":r["title"],"Rule":r["aggregation"],"Departments":g.Department.nunique(),"Records":len(g),"Status":"PENDING IQAC APPROVAL"})
    t1.dataframe(pd.DataFrame(props),use_container_width=True)
    deps=sorted(inv["Department"].dropna().unique()) if len(inv) else []
    ev=[]
    for m,req in NT.items():
        for dd in deps:
            g=raw[(raw["AQAR Metric"]==m)&(raw.Department==dd)]
            ev.append({"Department":dd,"Metric":m,"Required Information":req,"Submitted":"YES" if len(g) else "NO","Files":"; ".join(g["Relative Path"].unique()) if len(g) else "","IQAC Status":"VERIFY" if len(g) else "MISSING - ACTION REQUIRED"})
    t2.dataframe(pd.DataFrame(ev),use_container_width=True)
    with t3.expander("Evidence mapping"):st.dataframe(raw,use_container_width=True)
    with t3.expander("Errors"):st.dataframe(err if len(err) else pd.DataFrame([{"Status":"No errors"}]),use_container_width=True)
    st.download_button("⬇️ Download AQAR Master Excel",workbook(root,inv,raw,err).getvalue(),"AQAR_Master_Consolidation_Streamlit.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
